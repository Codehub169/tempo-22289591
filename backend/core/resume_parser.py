import os
import re
import spacy
from docx import Document
from pdfminer.high_level import extract_text

# Load the spaCy model. Make sure to download it first: python -m spacy download en_core_web_sm
try:
    nlp = spacy.load('en_core_web_sm')
except OSError:
    # This is a fallback for environments where the model might not be downloaded by startup.sh yet
    # or if spacy.cli.download is not available/working. Production setup should ensure model is present.
    print("spaCy model 'en_core_web_sm' not found. Please run 'python -m spacy download en_core_web_sm'")
    # A dummy nlp object to avoid crashing, though NER will not work.
    class DummyNLP:
        def __call__(self, text):
            class DummyDoc:
                def __init__(self, text):
                    self.text = text
                    self.ents = []
            return DummyDoc(text)
    nlp = DummyNLP()

# Regex patterns for contact information
EMAIL_REGEX = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
PHONE_REGEX = r"(?:\+?\d{1,3}[-\s\(\)]?)?(?:\d{2,4}[-\s\(\)]?){2,}\d{3,4}" # Simplified phone regex
LINKEDIN_REGEX = r"linkedin\.com/in/(\w[-_\w]*)"
GITHUB_REGEX = r"github\.com/(\w[-_\w]*)"
WEBSITE_REGEX = r"https?://(?:www\.)?([a-zA-Z0-9-]+(?:\.[a-zA-Z]+)+)(?:/[^\s]*)?"

# Keywords for sections
EXPERIENCE_KEYWORDS = ['experience', 'work history', 'employment history', 'professional experience']
EDUCATION_KEYWORDS = ['education', 'academic background', 'qualifications']
SKILLS_KEYWORDS = ['skills', 'technical skills', 'proficiencies', 'core competencies']
PROJECTS_KEYWORDS = ['projects', 'personal projects', 'portfolio']
SUMMARY_KEYWORDS = ['summary', 'profile', 'about me', 'objective']

def _extract_text_from_pdf(file_path):
    """Extracts text content from a PDF file."""
    try:
        return extract_text(file_path)
    except Exception as e:
        print(f"Error extracting text from PDF {file_path}: {e}")
        return ""

def _extract_text_from_docx(file_path):
    """Extracts text content from a DOCX file."""
    try:
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        print(f"Error extracting text from DOCX {file_path}: {e}")
        return ""

def _extract_name(text):
    """Extracts a person's name using spaCy NER."""
    doc = nlp(text)
    for ent in doc.ents:
        if ent.label_ == 'PERSON':
            # Take the first PERSON entity found, often at the beginning
            return ent.text
    # Fallback: look for capitalized words at the beginning if PERSON not found
    lines = text.split('\n')
    if lines:
        potential_name_line = lines[0].strip()
        # Simple check: if it's 2-3 capitalized words
        if re.match(r"^([A-Z][a-z]+(?:\s[A-Z][a-z]+){1,2})$", potential_name_line):
            return potential_name_line
    return "Your Name" # Default

def _extract_contact_info(text):
    """Extracts contact information using regex."""
    emails = list(set(re.findall(EMAIL_REGEX, text, re.IGNORECASE)))
    phones = list(set(re.findall(PHONE_REGEX, text)))
    
    linkedin_user = None
    linkedin_match = re.search(LINKEDIN_REGEX, text, re.IGNORECASE)
    if linkedin_match and linkedin_match.group(1):
        linkedin_user = f"linkedin.com/in/{linkedin_match.group(1)}"

    github_user = None
    github_match = re.search(GITHUB_REGEX, text, re.IGNORECASE)
    if github_match and github_match.group(1):
        github_user = f"github.com/{github_match.group(1)}"

    websites = list(set(re.findall(WEBSITE_REGEX, text, re.IGNORECASE)))
    # Filter out common social media domains if they are not the primary website
    generic_social = ['linkedin.com', 'github.com', 'twitter.com', 'facebook.com']
    personal_website = next((w[0] for w in websites if not any(social in w[0] for social in generic_social)), None)
    if personal_website and not personal_website.startswith(('http', 'https')):
        personal_website = 'http://' + personal_website

    return {
        "email": emails[0] if emails else "your.email@example.com",
        "phone": phones[0] if phones else "+1234567890",
        "linkedin": linkedin_user or "your-linkedin-profile",
        "github": github_user or "your-github-profile",
        "website": personal_website or "yourpersonal.website"
    }

def _find_section_content(text, keywords, next_section_keywords_lists=None):
    """Rudimentary section content extraction based on keywords."""
    text_lower = text.lower()
    start_index = -1
    
    for keyword in keywords:
        match = re.search(r"(^|\n)\s*" + re.escape(keyword) + r"\s*($|\n|:)", text_lower, re.MULTILINE)
        if match:
            start_index = match.end()
            break
    
    if start_index == -1:
        return ""

    end_index = len(text)
    if next_section_keywords_lists:
        for kw_list in next_section_keywords_lists:
            for kw in kw_list:
                # Search for the start of the next section
                next_match = re.search(r"(^|\n)\s*" + re.escape(kw) + r"\s*($|\n|:)", text_lower[start_index:], re.MULTILINE)
                if next_match:
                    current_end_index = start_index + next_match.start()
                    if current_end_index < end_index:
                        end_index = current_end_index
                        
    section_text = text[start_index:end_index].strip()
    # Clean up common resume section formatting (e.g., remove the title itself if it was included)
    # This is tricky and highly dependent on resume format.
    # For now, just return the block.
    return section_text

def _parse_experience(text_block):
    """Basic parsing of experience entries. Assumes chronological or distinct entries."""
    # This is highly heuristic. A more robust parser would use NLP to identify job titles, companies, dates.
    # For MVP, we'll split by common delimiters like multiple newlines or lines that look like date ranges / company names.
    # This is a placeholder for more sophisticated parsing logic.
    if not text_block.strip(): return []
    # Simplistic: assume each paragraph or double-newline separated block is an entry
    entries = re.split(r'\n\s*\n', text_block.strip()) # Split by blank lines
    parsed_entries = []
    for entry in entries:
        if not entry.strip(): continue
        lines = entry.split('\n')
        title = lines[0].strip() if lines else "Job Title"
        company_dates = lines[1].strip() if len(lines) > 1 else "Company & Dates"
        description = "\n".join(lines[2:]).strip() if len(lines) > 2 else "Description of responsibilities."
        parsed_entries.append({
            "title": title,
            "company": company_dates, # Needs further splitting for company and dates ideally
            "dates": "", # Placeholder, as splitting company_dates is complex
            "description": description
        })
    if not parsed_entries and text_block: # Fallback if split didn't work
        return [{'title': 'Experience Details', 'company': '', 'dates': '', 'description': text_block}]
    return parsed_entries

def _parse_education(text_block):
    if not text_block.strip(): return []
    entries = re.split(r'\n\s*\n', text_block.strip()) 
    parsed_entries = []
    for entry in entries:
        if not entry.strip(): continue
        lines = entry.split('\n')
        degree = lines[0].strip() if lines else "Degree"
        institution_dates = lines[1].strip() if len(lines) > 1 else "Institution & Dates"
        details = "\n".join(lines[2:]).strip() if len(lines) > 2 else "Details about education."
        parsed_entries.append({
            "degree": degree,
            "institution": institution_dates, # Needs further splitting
            "dates": "", # Placeholder
            "details": details
        })
    if not parsed_entries and text_block: # Fallback
        return [{'degree': 'Education Details', 'institution': '', 'dates': '', 'details': text_block}]
    return parsed_entries

def _parse_skills(text_block):
    if not text_block.strip(): return []
    # Skills are often comma-separated, bullet points, or lines
    skills = re.split(r'[\n,;]|\s{2,}', text_block) # Split by newline, comma, semicolon, or multiple spaces
    return [skill.strip() for skill in skills if skill.strip() and len(skill.strip()) > 1]

def _parse_projects(text_block):
    if not text_block.strip(): return []
    entries = re.split(r'\n\s*\n', text_block.strip())
    parsed_entries = []
    for entry in entries:
        if not entry.strip(): continue
        lines = entry.split('\n')
        name = lines[0].strip() if lines else "Project Name"
        description_tech = "\n".join(lines[1:]).strip() if len(lines) > 1 else "Description and technologies."
        parsed_entries.append({
            "name": name,
            "description": description_tech, # Could be split further
            "technologies": [], # Placeholder
            "link": ""
        })
    if not parsed_entries and text_block: # Fallback
        return [{'name': 'Project Details', 'description': text_block, 'technologies':[], 'link': ''}]
    return parsed_entries

def parse_resume(file_path):
    """Parses a resume file (PDF or DOCX) and extracts structured information."""
    _, extension = os.path.splitext(file_path)
    raw_text = ""

    if extension.lower() == '.pdf':
        raw_text = _extract_text_from_pdf(file_path)
    elif extension.lower() == '.docx':
        raw_text = _extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file type. Only PDF and DOCX are supported.")

    if not raw_text.strip():
        return {"error": "Could not extract text from resume."}

    # Basic data structure
    parsed_data = {
        "name": "Your Name",
        "title": "Professional Title", # Default, can be hard to extract
        "email": "your.email@example.com",
        "phone": "",
        "linkedin": "",
        "github": "",
        "website": "",
        "summary": "",
        "experience": [],
        "education": [],
        "skills": [],
        "projects": [],
        "profile_image_url": "https://images.unsplash.com/photo-1507003211169-0a1dd7228f2d?ixlib=rb-4.0.3&ixid=M3wxMjA3fDB8MHxwaG90by1wYWdlfHx8fGVufDB8fHx8fA%3D%3D&auto=format&fit=crop&w=387&q=80"
    }

    # Attempt to extract name from the top part of the resume
    # Consider the first few lines for name extraction to improve accuracy
    name_candidate_text = "\n".join(raw_text.split('\n')[:5]) 
    parsed_data["name"] = _extract_name(name_candidate_text)

    # Extract contact info from the whole text
    contact_info = _extract_contact_info(raw_text)
    parsed_data.update(contact_info)

    # Define section keyword lists for ordered extraction
    all_section_keywords = [
        SUMMARY_KEYWORDS, EXPERIENCE_KEYWORDS, EDUCATION_KEYWORDS, SKILLS_KEYWORDS, PROJECTS_KEYWORDS
    ]

    # Extract Summary
    summary_text = _find_section_content(raw_text, SUMMARY_KEYWORDS, [EXPERIENCE_KEYWORDS, EDUCATION_KEYWORDS, SKILLS_KEYWORDS, PROJECTS_KEYWORDS])
    parsed_data["summary"] = summary_text.strip() if summary_text else "A brief professional summary about yourself."

    # Extract Experience
    experience_text = _find_section_content(raw_text, EXPERIENCE_KEYWORDS, [EDUCATION_KEYWORDS, SKILLS_KEYWORDS, PROJECTS_KEYWORDS])
    parsed_data["experience"] = _parse_experience(experience_text)

    # Extract Education
    education_text = _find_section_content(raw_text, EDUCATION_KEYWORDS, [SKILLS_KEYWORDS, PROJECTS_KEYWORDS, EXPERIENCE_KEYWORDS]) # Experience might appear after education too
    parsed_data["education"] = _parse_education(education_text)

    # Extract Skills
    skills_text = _find_section_content(raw_text, SKILLS_KEYWORDS, [PROJECTS_KEYWORDS, EXPERIENCE_KEYWORDS, EDUCATION_KEYWORDS])
    parsed_data["skills"] = _parse_skills(skills_text)

    # Extract Projects
    projects_text = _find_section_content(raw_text, PROJECTS_KEYWORDS, [EXPERIENCE_KEYWORDS, EDUCATION_KEYWORDS, SKILLS_KEYWORDS]) # Check against other sections
    parsed_data["projects"] = _parse_projects(projects_text)
    
    # A very basic attempt to get a 'title' (e.g., Software Engineer)
    # This could be the first line of the summary, or the first job title.
    if parsed_data["experience"] and parsed_data["experience"][0]["title"] != "Job Title":
        parsed_data["title"] = parsed_data["experience"][0]["title"]
    elif parsed_data["summary"]:
        first_summary_line = parsed_data["summary"].split('\n')[0].strip()
        # If the first line is short and seems like a title
        if len(first_summary_line.split()) < 5 and first_summary_line:
            parsed_data["title"] = first_summary_line

    return parsed_data

# Example usage (for testing locally)
# if __name__ == '__main__':
#     # Create dummy files for testing
#     if not os.path.exists('dummy.pdf'): open('dummy.pdf', 'w').write('This is a dummy PDF.')
#     if not os.path.exists('dummy.docx'): Document().save('dummy.docx')

#     # Test with a dummy PDF (replace with a real PDF path for actual testing)
#     # data_pdf = parse_resume('path_to_your_resume.pdf') 
#     # print("PDF Data:", data_pdf)

#     # Test with a dummy DOCX (replace with a real DOCX path for actual testing)
#     # data_docx = parse_resume('path_to_your_resume.docx')
#     # print("DOCX Data:", data_docx)
